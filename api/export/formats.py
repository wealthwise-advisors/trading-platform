"""Shared multi-format export helpers -- turn one or more labeled DataFrames
into CSV, Excel, PDF, or Word bytes. Used by both the raw OHLC data export
endpoint and the backtest report export (metrics + trade log).
"""

from __future__ import annotations

import io

import pandas as pd

_MAX_TABLE_ROWS = 2000  # PDF/Word render row-by-row -- cap so a huge trade
                        # log or a wide date range doesn't take forever/blow up.

_MEDIA_TYPES = {
    "csv": "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def media_type_for(fmt: str) -> str:
    return _MEDIA_TYPES.get(fmt, "application/octet-stream")


def csv_bytes(df: pd.DataFrame) -> bytes:
    buf = io.StringIO()
    df.to_csv(buf)
    return buf.getvalue().encode("utf-8")


def excel_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    """sheets: {sheet_name: dataframe}. Sheet names are truncated to Excel's
    31-char limit."""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        for name, df in sheets.items():
            df.to_excel(writer, sheet_name=name[:31])
    return buf.getvalue()


def pdf_bytes(title: str, sections: list[tuple[str, pd.DataFrame]], subtitle: str = "") -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter),
                            topMargin=0.5 * inch, bottomMargin=0.5 * inch,
                            leftMargin=0.5 * inch, rightMargin=0.5 * inch)
    styles = getSampleStyleSheet()
    elements = [Paragraph(title, styles["Title"])]
    if subtitle:
        elements.append(Paragraph(subtitle, styles["Normal"]))
    elements.append(Spacer(1, 14))

    for heading, df in sections:
        elements.append(Paragraph(heading, styles["Heading2"]))
        elements.append(Spacer(1, 6))

        truncated = len(df) > _MAX_TABLE_ROWS
        show_df = df.head(_MAX_TABLE_ROWS).reset_index()
        data = [[str(c) for c in show_df.columns]] + show_df.astype(str).values.tolist()

        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#21262d")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fa")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elements.append(table)
        if truncated:
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(
                f"Showing first {_MAX_TABLE_ROWS} of {len(df)} rows.", styles["Italic"]))
        elements.append(Spacer(1, 18))

    doc.build(elements)
    return buf.getvalue()


def docx_bytes(title: str, sections: list[tuple[str, pd.DataFrame]], subtitle: str = "") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    if subtitle:
        doc.add_paragraph(subtitle)

    for heading, df in sections:
        doc.add_heading(heading, level=2)

        truncated = len(df) > _MAX_TABLE_ROWS
        show_df = df.head(_MAX_TABLE_ROWS).reset_index()

        table = doc.add_table(rows=1, cols=len(show_df.columns))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(show_df.columns):
            hdr_cells[i].text = str(col)
        for _, row in show_df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

        if truncated:
            doc.add_paragraph(f"Showing first {_MAX_TABLE_ROWS} of {len(df)} rows.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
